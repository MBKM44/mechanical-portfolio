import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips

SKILL_DOCS = Path("C:/Users/Mbkm44/.codex/plugins/cache/openai-primary-runtime/documents/26.601.10930/skills/documents")
sys.path.insert(0, str(SKILL_DOCS / "scripts"))
from table_geometry import apply_table_geometry  # noqa: E402

OUTPUT_DIR = Path("C:/AI_Codex/mechanical-portfolio/08_UAE_Mechanical_Job_Market_Research")
DATA_PATH = OUTPUT_DIR / "build" / "research_data.json"
DOCX_PATH = OUTPUT_DIR / "UAE_Mechanical_Engineering_Portfolio_Strategy_for_Mubarak_Al_Hamadi.docx"

ACCENT = "2E74B5"
DARK = "1F4D78"
HEADER_FILL = "F2F4F7"
BORDER = "D9D9D9"
BODY = "222222"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color=BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_run_font(run, size=11, bold=False, color=BODY):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc, text="", style=None, bold_lead=None):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.1
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = paragraph.add_run(text[len(bold_lead):])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.1
        run = paragraph.add_run(item)
        set_run_font(run)


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    if level == 1:
        paragraph.paragraph_format.space_before = Pt(16)
        paragraph.paragraph_format.space_after = Pt(8)
        size = 16
        color = ACCENT
    elif level == 2:
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)
        size = 13
        color = ACCENT
    else:
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(4)
        size = 12
        color = DARK
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=True, color=color)
    return paragraph


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BODY)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, ACCENT, 16, 8),
        ("Heading 2", 13, ACCENT, 12, 6),
        ("Heading 3", 12, DARK, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1

    for name in ["List Bullet", "List Number"]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.167


def set_cell_text(cell, text, bold=False, size=9, color=BODY):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths_dxa, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    table.allow_autofit = False
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, header, bold=True, size=font_size, color="000000")
        set_cell_shading(cell, HEADER_FILL)
        set_cell_borders(cell)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            set_cell_text(cell, value, size=font_size)
            set_cell_borders(cell)
    apply_table_geometry(table, widths_dxa, table_width_dxa=9360, indent_dxa=120)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(4)
    return table


def add_label_table(doc, pairs):
    rows = [[label, value] for label, value in pairs]
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True, size=9, color=DARK)
        set_cell_shading(cells[0], HEADER_FILL)
        set_cell_borders(cells[0])
        set_cell_text(cells[1], value, size=9)
        set_cell_borders(cells[1])
    apply_table_geometry(table, [2100, 7260], table_width_dxa=9360, indent_dxa=120)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)
    return table


def join_sources(*ids):
    return ", ".join(ids)


def source_index(sources):
    return {s["id"]: s for s in sources}


def build_doc():
    data = json.loads(DATA_PATH.read_text(encoding="utf-8"))
    doc = Document()
    configure_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("UAE Mechanical Engineering Portfolio Strategy for Mubarak Al Hamadi")
    set_run_font(run, size=22, bold=True, color="0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run(f"Prepared from UAE job-market research and current portfolio/CV evidence | {data['generatedDate']}")
    set_run_font(run, size=10, color="555555")

    add_heading(doc, "Executive Summary", 1)
    add_paragraph(
        doc,
        "The strongest portfolio direction is not to add random mechanical projects. It is to convert your existing AWH, compressor, CAD, simulation, and mechatronics evidence into the work products UAE employers repeatedly request: calculations, drawings, P&IDs, QA/QC documents, maintenance plans, dashboards, and short technical reports."
    )
    add_paragraph(
        doc,
        "Start first with a district cooling ETS and chilled-water calculation pack. It is UAE-specific, directly connected to your refrigeration and energy-analysis background, low cost to execute, and useful for MEP, district cooling, application engineering, facilities, and energy-system interviews."
    )

    add_heading(doc, "Market Research Findings", 1)
    add_bullets(doc, data["marketFindings"])

    add_heading(doc, "Main Role Clusters Suitable for You", 1)
    role_rows = []
    for r in data["roleSummary"]:
        role_rows.append([
            r["roleCluster"],
            r["requirements"],
            r["relevance"],
            f"{r['priority']} / {r['difficulty']}"
        ])
    add_table(
        doc,
        ["Role cluster", "Repeated market requirements", "Fit to CV", "Priority / difficulty"],
        role_rows,
        [1900, 3900, 2400, 1160],
        font_size=8
    )

    add_heading(doc, "Skills Employers Repeatedly Request", 1)
    top_skills = data["skillsGap"][:12]
    skill_rows = [[s["requiredSkill"], s["marketFrequency"], s["gapLevel"], s["proofProject"]] for s in top_skills]
    add_table(
        doc,
        ["Required skill", "Market frequency", "Gap", "Best proof to create"],
        skill_rows,
        [2200, 1500, 900, 4760],
        font_size=8.5
    )

    add_heading(doc, "Your CV Strengths", 1)
    add_bullets(doc, data["cvStrengths"])

    add_heading(doc, "Your CV Gaps", 1)
    add_bullets(doc, data["cvGaps"])

    add_heading(doc, "Recommended Portfolio Direction", 1)
    add_paragraph(
        doc,
        "Use your portfolio to show engineering work products, not only project descriptions. Each new case study should have a visible calculation, drawing, system diagram, checklist, dashboard, or report that a hiring manager can inspect in under two minutes."
    )
    add_table(
        doc,
        ["Direction", "What to produce", "Why it helps"],
        [
            ["MEP/district cooling bridge", "ETS, chilled water pump, HVAC load, heat exchanger, and Revit/AutoCAD deliverables", "Covers the highest-volume UAE entry route and connects naturally to your AWH/refrigeration work."],
            ["Maintenance/reliability bridge", "Compressor RCA, PM plan, FMEA, CMMS dashboard, safety checklist", "Turns Atlas Copco experience into visible proof for industrial, compressor, and oil/gas O&M roles."],
            ["Manufacturing/QA bridge", "GD&T drawing, BOM, DFM note, ITP, inspection checklist, hydrotest checklist", "Makes your CAD/prototype experience relevant to manufacturing, defense, QA/QC, and site roles."]
        ],
        [2200, 3700, 3460],
        font_size=8.5
    )

    add_heading(doc, "Best 5 Portfolio Projects", 1)
    for project in data["bestProjects"]:
        add_heading(doc, project["project"], 2)
        add_label_table(doc, [
            ("Problem definition", project["problem"]),
            ("Why UAE market cares", project["whyMarketCares"]),
            ("What to design/build/analyze", project["whatToBuild"]),
            ("Minimum viable version", project["mvp"]),
            ("Advanced version", project["advanced"]),
            ("Exact deliverables", project["deliverables"]),
            ("What to show in portfolio", project["portfolio"]),
            ("What to say in interview", project["interview"]),
            ("Execution steps", project["steps"])
        ])

    add_heading(doc, "Exact Project Execution Steps", 1)
    add_paragraph(
        doc,
        "Execution order: finish one project to portfolio-ready quality before opening too many parallel threads. The first two projects should become complete case studies; the next three can begin as MVP artifact packs."
    )
    plan_rows = [[p["week"], p["focus"], p["tasks"], p["outputs"], p["acceptance"]] for p in data["plan30"]]
    add_table(
        doc,
        ["Week", "Focus", "Tasks", "Outputs", "Acceptance check"],
        plan_rows,
        [900, 1700, 3300, 1900, 1560],
        font_size=8.3
    )

    add_heading(doc, "Suggested Portfolio Wording", 1)
    wording_rows = [[p["project"].split(". ", 1)[1], p["portfolioWording"], p["interview"]] for p in data["bestProjects"]]
    add_table(
        doc,
        ["Project", "Portfolio wording", "Interview angle"],
        wording_rows,
        [1900, 3800, 3660],
        font_size=8.3
    )

    add_heading(doc, "Suggested LinkedIn/CV Bullet Points", 1)
    bullet_rows = []
    for p in data["bestProjects"]:
        bullet_rows.append([p["project"].split(". ", 1)[1], "\n".join(p["bullets"])])
    add_table(
        doc,
        ["Project", "Bullet points after completion"],
        bullet_rows,
        [2300, 7060],
        font_size=8.5
    )

    add_heading(doc, "Portfolio Project Idea Bank", 1)
    idea_rows = []
    for item in data["projectIdeas"][:15]:
        idea_rows.append([
            str(item["priorityRanking"]),
            item["projectTitle"],
            item["targetRole"],
            item["deliverables"],
            item["impact"]
        ])
    add_table(
        doc,
        ["Rank", "Project", "Target role", "Core deliverables", "Impact"],
        idea_rows,
        [600, 2400, 1900, 3360, 1100],
        font_size=8
    )

    add_heading(doc, "Evidence Limits and Uncertainty", 1)
    add_bullets(doc, [
        "Job posts can change or expire quickly; this research uses sources available on June 7, 2026.",
        "LinkedIn and some employer pages are dynamic or partially login-gated, so LinkedIn evidence is treated as medium confidence unless supported by another source.",
        "Market frequency is a qualitative ranking from repeated requirements across live job boards and employer pages, not a complete scrape of every UAE mechanical job.",
        "The recommendations intentionally avoid claiming professional authority in regulated design areas. Portfolio projects should be framed as concept or calculation packs unless reviewed by a licensed specialist."
    ])

    add_heading(doc, "Sources and Links", 1)
    source_rows = [[s["id"], s["source"], s["evidence"], s["url"]] for s in data["sources"]]
    add_table(
        doc,
        ["ID", "Source", "Evidence used", "URL"],
        source_rows,
        [550, 2200, 3800, 2810],
        font_size=7.2
    )

    add_heading(doc, "Final Recommendation", 1)
    add_paragraph(
        doc,
        "Start with the District Cooling ETS Design and Chilled Water Calculation Pack. It is the best first project because it is UAE-specific, low cost, fast to execute, directly connected to your AWH refrigeration and energy-analysis story, and valuable across MEP, district cooling, HVAC application engineering, facilities, and energy roles."
    )
    add_paragraph(
        doc,
        "The second project should be the Compressor Maintenance, RCA, and Reliability Dashboard. Together, these two projects will make your portfolio read as both UAE-building-services aware and industrial-maintenance aware, which is exactly the bridge your CV needs."
    )

    # Quiet footer.
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("UAE Mechanical Engineering Portfolio Strategy")
    set_run_font(footer_run, size=8, color="777777")

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
    print(DOCX_PATH)
